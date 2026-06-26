"""Extracción estructurada de requisitos de pliegos SECOP II.

Este módulo analiza el texto del pliego y sus anexos (Matriz 1 Experiencia,
Matriz 2 Indicadores, etc.) para extraer requisitos cuantificables:
- Experiencia mínima (contratos, valor, tipo de obra).
- Capacidad financiera (patrimonio, indicadores).
- Capacidad residual.
- Documentos habilitantes y de oferta.
- Factores de calidad / puntaje.

Es un extractor por reglas/heurísticas, diseñado para procesos de
infraestructura de transporte (Documentos Tipo CCE-EICP).
"""

import json
import re
from pathlib import Path
from typing import Any

from openpyxl import load_workbook
from docx import Document


SMMLV = 1_423_500  # Valor por defecto; se puede parametrizar vía env.


def _normalizar(texto: str) -> str:
    """Normaliza texto para matching robusto."""
    t = texto.lower()
    t = t.replace("á", "a").replace("é", "e").replace("í", "i")
    t = t.replace("ó", "o").replace("ú", "u").replace("ñ", "n")
    t = re.sub(r"[^a-z0-9$.,:/\-%=\s]+", " ", t)
    t = re.sub(r"\s+", " ", t)
    return t.strip()


def _buscar_valor_porcentaje(texto: str, contextos: list[str], palabras_excluir: list[str] | None = None) -> list[dict]:
    """Busca porcentajes cercanos a ciertos contextos (palabras clave).

    Filtra ventanas que contengan palabras excluidas (ej. anticipo, AIU) para
    evitar falsos positivos.
    """
    resultados = []
    texto_norm = _normalizar(texto)
    excluir = [_normalizar(p) for p in (palabras_excluir or [])]
    for ctx in contextos:
        for m in re.finditer(re.escape(_normalizar(ctx)), texto_norm):
            inicio = max(0, m.start() - 200)
            fin = min(len(texto_norm), m.end() + 400)
            ventana = texto_norm[inicio:fin]
            if any(exc in ventana for exc in excluir):
                continue
            for match in re.finditer(r"(\d{1,3})(?:\.\d+)?\s*%", ventana):
                resultados.append({
                    "contexto": ctx,
                    "valor": float(match.group(1)),
                    "unidad": "%",
                    "texto": ventana,
                })
    return resultados


def _buscar_numeros_con_unidad(texto: str, unidad: str, contextos: list[str]) -> list[dict]:
    """Busca números seguidos de una unidad (SMMLV, contratos, etc.)."""
    resultados = []
    texto_norm = _normalizar(texto)
    unidad_norm = _normalizar(unidad)
    for ctx in contextos:
        for m in re.finditer(re.escape(_normalizar(ctx)), texto_norm):
            inicio = max(0, m.start() - 150)
            fin = min(len(texto_norm), m.end() + 300)
            ventana = texto_norm[inicio:fin]
            # número seguido de la unidad, o unidad seguida de número
            patron = rf"(\d{{1,6}})(?:\.\d+)?\s*{unidad_norm}|{unidad_norm}\s*(\d{{1,6}})(?:\.\d+)?"
            for match in re.finditer(patron, ventana):
                valor = match.group(1) or match.group(2)
                if valor:
                    resultados.append({
                        "contexto": ctx,
                        "valor": float(valor),
                        "unidad": unidad,
                        "texto": ventana,
                    })
    return resultados


def _extraer_actividad_matriz1(path: str, actividad_buscada: str) -> dict | None:
    """Busca una actividad en la Matriz 1 — Experiencia (Excel) y retorna sus requisitos."""
    if not Path(path).exists():
        return None

    try:
        wb = load_workbook(path, data_only=True)
    except Exception:
        return None

    sheet = wb.worksheets[0]
    actividad_norm = _normalizar(actividad_buscada)

    for i, row in enumerate(sheet.iter_rows(values_only=True), 1):
        celdas = [str(c) if c is not None else "" for c in row]
        fila_texto = " ".join(celdas)
        if actividad_norm in _normalizar(fila_texto):
            # La fila de la actividad suele incluir GENERAL; las siguientes ESPECIFICA y DIMENSIONAMIENTO
            general = None
            especifica = None
            dimensionamiento = None
            try:
                for offset in range(0, 6):
                    next_row = [str(c.value) if c.value is not None else "" for c in sheet[i + offset]]
                    joined = " ".join(next_row).strip()
                    if not joined:
                        continue
                    joined_norm = _normalizar(joined)
                    # La primera fila con GENERAL la tomamos como experiencia general
                    if "general" in joined_norm and not general:
                        general = joined
                    elif ("especifica" in joined_norm or "específica" in joined_norm) and not especifica:
                        especifica = joined
                    elif "dimensionamiento" in joined_norm and not dimensionamiento:
                        dimensionamiento = joined
            except Exception:
                pass

            return {
                "actividad": actividad_buscada,
                "fila": i,
                "experiencia_general": general,
                "experiencia_especifica": especifica,
                "dimensionamiento": dimensionamiento,
            }

    return None


def _encontrar_matriz1(documentos: list[Any]) -> str | None:
    """Busca el archivo de Matriz 1 Experiencia entre los documentos descargados."""
    candidatos = []
    for doc in documentos:
        nombre = (doc.nombre or doc.filename or "").lower()
        path = doc.path or ""
        if not path or not Path(path).exists():
            continue
        # Puntuación de matching
        score = 0
        if "matriz1" in nombre or "matriz 1" in nombre:
            score += 3
        if "matriz1 experiencia" in nombre or "matriz 1 experiencia" in nombre:
            score += 2
        if "experiencia" in nombre and ("matriz1" in nombre or "matriz 1" in nombre):
            score += 1
        # Excluir archivos que claramente son otra matriz
        if any(x in nombre for x in ["indicador", "riesgo", "bienes relevantes", "bienes_relevantes"]):
            score -= 5
        if score > 0:
            candidatos.append((score, path))
    if candidatos:
        candidatos.sort(reverse=True)
        return candidatos[0][1]
    return None


def _encontrar_matriz2(documentos: list[Any]) -> str | None:
    """Busca el archivo de Matriz 2 Indicadores Financieros entre los documentos descargados."""
    candidatos = []
    for doc in documentos:
        nombre = (doc.nombre or doc.filename or "").lower()
        path = doc.path or ""
        if not path or not Path(path).exists():
            continue
        score = 0
        if "matriz2" in nombre or "matriz 2" in nombre:
            score += 3
        if "indicadores financieros" in nombre or "indicadores y organizacionales" in nombre:
            score += 2
        if "matriz" in nombre and "indicador" in nombre:
            score += 1
        # Excluir otras matrices
        if any(x in nombre for x in ["experiencia", "riesgo", "bienes relevantes", "bienes_relevantes"]):
            score -= 5
        if score > 0:
            candidatos.append((score, path))
    if candidatos:
        candidatos.sort(reverse=True)
        return candidatos[0][1]
    return None


def _extraer_valor_minimo_smmlv(texto_pliego: str) -> dict[str, Any] | None:
    """Extrae el valor mínimo de experiencia expresado en SMMLV del pliego.

    Busca patrones como:
        - "SMMLV: 282.14"
        - "... expresado en SMMLV" cerca de un número
        - "valor mínimo a certificar ... 282.14 SMMLV"
    Retorna {"smmlv": float, "cop": float} o None.
    """
    texto_norm = _normalizar(texto_pliego)

    # Patrón directo: "SMMLV: 282.14" o "SMMLV 282.14"
    patron_directo = re.search(
        r"smmlv[:\s]+(\d{1,3}(?:[.,]\d+)?)",
        texto_norm,
    )
    if patron_directo:
        valor_str = patron_directo.group(1).replace(",", ".")
        try:
            smmlv = float(valor_str)
            return {"smmlv": smmlv, "cop": round(smmlv * SMMLV), "fuente": "SMMLV directo"}
        except ValueError:
            pass

    # Patrón: "valor mínimo a certificar ... N SMMLV"
    for m in re.finditer(r"valor\s*minimo\s*a\s*certificar", texto_norm):
        ventana = texto_norm[max(0, m.start() - 100):min(len(texto_norm), m.end() + 200)]
        match = re.search(r"(\d{1,3}(?:[.,]\d+)?)\s*smmlv", ventana)
        if match:
            try:
                smmlv = float(match.group(1).replace(",", "."))
                return {"smmlv": smmlv, "cop": round(smmlv * SMMLV), "fuente": "valor mínimo a certificar"}
            except ValueError:
                pass

    return None


def _extraer_presupuesto_oficial_pliego(texto_pliego: str) -> int | None:
    """Extrae el presupuesto oficial del cuadro 1.1 del pliego (pesos incluido IVA)."""
    texto_norm = _normalizar(texto_pliego)
    # Buscar "presupuesto oficial" seguido de un valor monetario en las siguientes 300 chars
    for m in re.finditer(r"presupuesto\s*oficial", texto_norm):
        ventana = texto_norm[max(0, m.start()):min(len(texto_norm), m.end() + 400)]
        match = re.search(r"\$\s*(\d{1,4}(?:[.,]\d{3})*(?:,\d{2})?)", ventana)
        if match:
            valor_str = match.group(1).replace(".", "").replace(",", ".")
            try:
                return int(float(valor_str))
            except ValueError:
                pass
    return None


def _extraer_matriz2(path: str) -> dict[str, Any] | None:
    """Extrae indicadores financieros y valores concertados de la Matriz 2 (DOCX).

    El documento típico contiene dos tablas: la primera para Mipyme y la segunda
    para los demás proponentes. Si solo hay una tabla se asume "general".
    """
    if not Path(path).exists():
        return None

    try:
        doc = Document(path)
    except Exception:
        return None

    # Determinar perfiles disponibles a partir de los párrafos que preceden a cada tabla.
    perfiles_por_tabla: list[str] = []
    tabla_idx = 0
    for para in doc.paragraphs:
        texto = para.text.lower()
        if "tabla" in texto or "indicador" in texto or "valor" in texto:
            if "mipyme" in texto and "demas" not in texto:
                perfiles_por_tabla.append("mipyme")
            elif "demas" in texto or "otros" in texto or "no mipyme" in texto:
                perfiles_por_tabla.append("general")
            else:
                perfiles_por_tabla.append("general")
            tabla_idx += 1

    indicadores: list[dict[str, Any]] = []
    tabla_count = 0
    for table in doc.tables:
        header = [cell.text.strip().lower() for cell in table.rows[0].cells] if table.rows else []
        if not header or "indicador" not in " ".join(header):
            continue

        # Identificar columna de indicador y valor.
        idx_indicador = next((i for i, h in enumerate(header) if "indicador" in h), 0)
        idx_valor = next((i for i, h in enumerate(header) if "valor" in h or "concertado" in h), 1)

        # Asignar perfil: si hay dos tablas, primera mipyme, segunda general.
        if len(perfiles_por_tabla) >= 2:
            perfil = perfiles_por_tabla[tabla_count] if tabla_count < len(perfiles_por_tabla) else "general"
        else:
            perfil = "general"
        tabla_count += 1

        for row in table.rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if not cells:
                continue
            texto_fila = " ".join(cells).lower()
            if any(p in texto_fila for p in ["mipyme", "demas", "otros", "no mipyme"]):
                continue

            nombre = cells[idx_indicador] if idx_indicador < len(cells) else ""
            valor = cells[idx_valor] if idx_valor < len(cells) else ""
            if not nombre:
                continue

            nombre_norm = _normalizar(nombre)
            if not nombre_norm or nombre_norm in ("indicador", "valor concertado"):
                continue

            indicadores.append({
                "perfil": perfil,
                "nombre": nombre,
                "valor_texto": valor,
                "valor_numerico": _extraer_numero_umbral(valor),
                "categoria": _categorizar_indicador(nombre_norm),
            })

    if not indicadores:
        return None

    return {
        "indicadores": indicadores,
        "resumen": _resumen_indicadores(indicadores),
        "categorias": sorted({i["categoria"] for i in indicadores}),
    }


def _extraer_numero_umbral(texto: str) -> float | None:
    """Extrae un número de un texto como '1,2', ',02', '1.300.000', 'Definido en pliegos'.

    Interpreta el formato colombiano: punto como separador de miles y coma como decimal.
    """
    if not texto:
        return None
    texto = texto.strip()
    # Casos como ",02" → "0.02"
    if texto.startswith(","):
        texto = "0" + texto
    # Eliminar separadores de miles (puntos entre dígitos) y convertir coma decimal a punto
    texto = re.sub(r"(?<=\d)\.(?=\d{3}(?:\D|$))", "", texto)
    texto = texto.replace(",", ".")
    match = re.search(r"(\d+(?:\.\d+)?)", texto)
    if match:
        try:
            return float(match.group(1))
        except ValueError:
            return None
    return None


def _categorizar_indicador(nombre_norm: str) -> str:
    """Clasifica un indicador financiero en una categoría estándar."""
    if "liquidez" in nombre_norm or "razon corriente" in nombre_norm:
        return "liquidez"
    if "endeudamiento" in nombre_norm:
        return "endeudamiento"
    if "cobertura" in nombre_norm or "intereses" in nombre_norm:
        return "cobertura"
    if "rentabilidad" in nombre_norm or "roe" in nombre_norm or "roi" in nombre_norm or "rentabilidad del patrimonio" in nombre_norm or "rentabilidad del activo" in nombre_norm:
        return "rentabilidad"
    if "capital de trabajo" in nombre_norm:
        return "capital_trabajo"
    return "otro"


def _resumen_indicadores(indicadores: list[dict[str, Any]]) -> dict[str, dict[str, Any]]:
    """Resume los indicadores por categoría y por perfil (mipyme / general).

    Toma el valor más restrictivo cuando hay duplicados: para endeudamiento
    el menor; para los demás el mayor.
    """
    resumen: dict[str, dict[str, Any]] = {}
    perfiles = sorted({i["perfil"] for i in indicadores})
    for perfil in perfiles:
        resumen[perfil] = {}
        for cat in ["liquidez", "endeudamiento", "cobertura", "rentabilidad", "capital_trabajo"]:
            vals = [
                i for i in indicadores
                if i["categoria"] == cat and i["perfil"] == perfil and i["valor_numerico"] is not None
            ]
            if vals:
                if cat == "endeudamiento":
                    seleccion = min(vals, key=lambda x: x["valor_numerico"])
                else:
                    seleccion = max(vals, key=lambda x: x["valor_numerico"])
                resumen[perfil][cat] = {
                    "valor_minimo": seleccion["valor_numerico"],
                    "texto": seleccion["valor_texto"],
                }
    return resumen


def _extraer_capacidad_residual(texto_pliego: str) -> dict[str, Any]:
    """Extrae del texto del pliego la metodología de capacidad residual.

    Retorna fórmulas de CRPC, requisito CRP >= CRPC y tabla de factores si se encuentran.
    """
    texto_norm = _normalizar(texto_pliego)
    resultado: dict[str, Any] = {"requerida": "capacidad residual" in texto_norm}

    if not resultado["requerida"]:
        return resultado

    # Fórmula CRPC para plazo <= 12 meses
    if re.search(r"crpc\s*=\s*poe\s*-?\s*anticipo", texto_norm):
        resultado["formula_crpc_corto_plazo"] = "CRPC = POE - Anticipo y/o pago anticipado"

    # Fórmula CRPC para plazo > 12 meses
    if re.search(
        r"poe\s*-?\s*anticipo[^\d]{0,50}plazo\s*estimado[^\d]{0,20}crpc\s*=\s*\d+",
        texto_norm,
    ):
        resultado["formula_crpc_largo_plazo"] = "CRPC = ((POE - Anticipo) / Plazo estimado meses) * 12"

    # Requisito CRP >= CRPC
    if re.search(
        r"capacidad residual del proponente.{0,150}(mayor|igual|superior).{0,150}capacidad residual del proceso",
        texto_norm,
    ):
        resultado["requisito_crp_crpc"] = "CRP >= CRPC"

    # Tabla de factores/puntajes (E, CF, CT). Se buscan puntajes de 2-3 dígitos
    # asociados explícitamente a cada factor para evitar ruido numérico.
    factores: list[dict[str, Any]] = []
    factores_regex = [
        ("experiencia", r"experiencia\s*\(?\s*e\s*\)?[^\d]{0,40}(\d{2,3})"),
        ("capacidad financiera", r"capacidad\s*financiera\s*\(?\s*cf\s*\)?[^\d]{0,40}(\d{2,3})"),
        ("capacidad tecnica", r"capacidad\s*tecnica\s*\(?\s*ct\s*\)?[^\d]{0,40}(\d{2,3})"),
    ]
    for nombre, patron in factores_regex:
        for m in re.finditer(patron, texto_norm):
            puntaje = int(m.group(1))
            codigo = nombre.split()[-1][0].upper() if len(nombre.split()) > 1 else nombre[0].upper()
            if codigo == "T":
                codigo = "CT"
            elif codigo == "F":
                codigo = "CF"
            factores.append({"nombre": nombre, "codigo": codigo, "puntaje_maximo": puntaje})
    if factores:
        # De-duplicar por código, quedarse con el puntaje más alto.
        factores_unicos: dict[str, dict[str, Any]] = {}
        for f in factores:
            codigo = f["codigo"]
            if codigo not in factores_unicos or f["puntaje_maximo"] > factores_unicos[codigo]["puntaje_maximo"]:
                factores_unicos[codigo] = f
        resultado["factores"] = list(factores_unicos.values())

    # Umbral porcentual explícito de CRP si existe
    pct_match = re.search(
        r"capacidad residual del proponente[^\n]{0,100}(mayor|igual|superior|inferior)[^\n]{0,100}(\d{1,3})(?:\.\d+)?\s*%",
        texto_norm,
    )
    if pct_match:
        resultado["min_crp_pct"] = float(pct_match.group(2))

    return resultado


def extraer_requisitos_estructurados(
    texto_pliego: str,
    documentos_proceso: list[Any],
    presupuesto: int,
    smmlv: int = SMMLV,
) -> dict[str, Any]:
    """Extrae requisitos cuantificables del pliego y sus anexos.

    Retorna un dict con:
        - tipo_proceso (infraestructura de transporte, etc.)
        - complejidad_tecnica
        - actividad_principal
        - experiencia: {min_contratos, max_contratos, tipo_obra, valor_minimo_po_pct, ...}
        - capacidad_financiera: {patrimonio_minimo_po_pct, indicadores, ...}
        - capacidad_residual: {...}
        - documentos_requeridos: list
        - factores_calidad: {...}
        - advertencias: list
    """
    texto_norm = _normalizar(texto_pliego)
    resultado: dict[str, Any] = {
        "tipo_proceso": None,
        "complejidad_tecnica": None,
        "actividad_principal": None,
        "experiencia": {},
        "capacidad_financiera": {},
        "capacidad_residual": {},
        "documentos_requeridos": [],
        "factores_calidad": {},
        "advertencias": [],
    }

    # 1. Tipo y complejidad del proceso
    if "infraestructura de transporte" in texto_norm:
        resultado["tipo_proceso"] = "Infraestructura de transporte"

    complejidad_match = re.search(
        r"complejidad\s*tecnica[\s\S]{0,120}?(baja[-\s]*media|baja|media|alta)",
        texto_norm,
    )
    if complejidad_match:
        resultado["complejidad_tecnica"] = complejidad_match.group(1).replace(" ", "-").strip("-")

    # 2. Actividad principal: intentar detectar actividades de infraestructura,
    # edificación, consultoría y otros sectores comunes en SECOP II.
    # Primero buscamos el formato del Documento Base: "requisitos de experiencia son:"
    # seguido de la sección y la actividad principal.
    actividad_detectada = None
    seccion_match = re.search(
        r"requisitos\s*de\s*experiencia\s*son[:\s]+"
        r"(\d+(?:\.\d+)?)\s+([a-z0-9\s\-]+?(?:"
        r"vias?\s+(?:primarias?|secundarias?|terciarias)|vias?\s+urbanas?|"
        r"puentes|tuneles|túneles|aeroportuarias?|ferreas?|férreas?|"
        r"edificacion|edificación|construccion|construcción|"
        r"consultoria|consultoría|diseno|diseño|interventoria|interventoría|"
        r"estudios|topografia|topografía|geotecnia|hidraulica|hidráulica|"
        r"mantenimiento|rehabilitacion|rehabilitación|modernizacion|modernización"
        r"))\s*\.?\s*"
        r"(\d+(?:\.\d+)?)\s+([a-z0-9\s\-]+?(?:"
        r"vias?\s+(?:primarias?|secundarias?|terciarias)|vias?\s+urbanas?|"
        r"puentes|tuneles|túneles|aeroportuarias?|ferreas?|férreas?|"
        r"edificacion|edificación|construccion|construcción|"
        r"consultoria|consultoría|diseno|diseño|interventoria|interventoría|"
        r"estudios|topografia|topografía|geotecnia|hidraulica|hidráulica|"
        r"mantenimiento|rehabilitacion|rehabilitación|modernizacion|modernización"
        r"))",
        texto_norm,
    )
    if seccion_match:
        actividad_detectada = {
            "seccion_codigo": seccion_match.group(1).strip(),
            "seccion_descripcion": seccion_match.group(2).strip().upper(),
            "codigo": seccion_match.group(3).strip(),
            "descripcion": seccion_match.group(4).strip().upper(),
        }
    else:
        # Fallback: buscar una sola actividad con código
        actividad_match = re.search(
            r"(\d+\.\d+)\s+([a-z0-9\s\-]+?(?:"
            r"vias?\s+terciarias|vias?\s+(primarias?|secundarias?)|vias?\s+urbanas?|"
            r"puentes|tuneles|túneles|aeroportuarias?|ferreas?|férreas?|"
            r"edificacion|edificación|construccion|construcción|"
            r"consultoria|consultoría|diseno|diseño|interventoria|interventoría|"
            r"estudios|topografia|topografía|geotecnia|hidraulica|hidráulica|"
            r"mantenimiento|rehabilitacion|rehabilitación|modernizacion|modernización"
            r"))\s*\.?",
            texto_norm,
        )
        if actividad_match:
            actividad_detectada = {
                "codigo": actividad_match.group(1).strip(),
                "descripcion": actividad_match.group(2).strip().upper(),
            }

    if actividad_detectada:
        resultado["actividad_principal"] = actividad_detectada

    # 3. Experiencia: número de contratos
    min_contratos = 1
    max_contratos = 5
    if re.search(r"minimo\s*uno\s*\(?1\)?\s*y\s*maximo\s*cinco\s*\(?5\)?", texto_norm):
        min_contratos = 1
        max_contratos = 5
    elif re.search(r"maximo\s*seis\s*\(?6\)?", texto_norm):
        max_contratos = 6
    elif re.search(r"maximo\s*siete\s*\(?7\)?", texto_norm):
        max_contratos = 7

    resultado["experiencia"]["min_contratos"] = min_contratos
    resultado["experiencia"]["max_contratos"] = max_contratos

    # 4. Tipo de obra requerida para experiencia general (basado en actividad principal)
    tipos_obra = []
    actividad_desc = (resultado.get("actividad_principal") or {}).get("descripcion", "")
    actividad_norm = _normalizar(actividad_desc)

    if "via terciaria" in actividad_norm or "vias terciarias" in actividad_norm:
        tipos_obra.append("Vías terciarias")
    if "puente" in actividad_norm:
        tipos_obra.append("Puentes")
    if "aeroportuario" in actividad_norm:
        tipos_obra.append("Obras aeroportuarias")
    if "tunel" in actividad_norm or "túnel" in actividad_norm:
        tipos_obra.append("Túneles")
    if "ferrea" in actividad_norm or "férrea" in actividad_norm:
        tipos_obra.append("Obras férreas")
    if "urbana" in actividad_norm:
        tipos_obra.append("Infraestructura vial urbana")
    if "primaria" in actividad_norm or "secundaria" in actividad_norm:
        tipos_obra.append("Vías primarias/secundarias")
    if "edificacion" in actividad_norm or "edificación" in actividad_norm:
        tipos_obra.append("Edificación")
    if "consultoria" in actividad_norm or "consultoría" in actividad_norm or "interventoria" in actividad_norm or "interventoría" in actividad_norm:
        tipos_obra.append("Servicios de ingeniería/consultoría")
    if "diseno" in actividad_norm or "diseño" in actividad_norm:
        tipos_obra.append("Diseño")
    if "topografia" in actividad_norm or "topografía" in actividad_norm:
        tipos_obra.append("Topografía")
    if "geotecnia" in actividad_norm:
        tipos_obra.append("Geotecnia")
    if "hidraulica" in actividad_norm or "hidráulica" in actividad_norm:
        tipos_obra.append("Obras hidráulicas")
    if "mantenimiento" in actividad_norm:
        tipos_obra.append("Mantenimiento vial")
    if "rehabilitacion" in actividad_norm or "rehabilitación" in actividad_norm or "modernizacion" in actividad_norm or "modernización" in actividad_norm:
        tipos_obra.append("Rehabilitación/modernización")
    if "pavimento asfaltico" in texto_norm or "concreto hidraulico" in texto_norm or "concreto hidráulico" in texto_norm:
        tipos_obra.append("Pavimento asfáltico/concreto hidráulico")
    if tipos_obra:
        resultado["experiencia"]["tipos_obra"] = tipos_obra

    # Presupuesto oficial según el pliego (puede diferir del valor de SECOP)
    presupuesto_pliego = _extraer_presupuesto_oficial_pliego(texto_pliego)
    if presupuesto_pliego:
        resultado["presupuesto_oficial_pliego"] = presupuesto_pliego
        # Si hay diferencia significativa con el presupuesto de SECOP, advertir
        if presupuesto > 0 and abs(presupuesto_pliego - presupuesto) / presupuesto > 0.05:
            resultado["advertencias"].append(
                f"El presupuesto del pliego (${presupuesto_pliego:,.0f} COP) difiere "
                f"del registrado en SECOP (${presupuesto:,.0f} COP)."
            )

    # 5. Valor mínimo de experiencia: primero intentar SMMLV explícito en el pliego
    min_smmlv = _extraer_valor_minimo_smmlv(texto_pliego)
    if min_smmlv:
        resultado["experiencia"]["valor_minimo_smmlv"] = min_smmlv["smmlv"]
        resultado["experiencia"]["valor_minimo_cop_smmlv"] = min_smmlv["cop"]
        resultado["experiencia"]["valor_minimo_cop"] = min_smmlv["cop"]
        resultado["experiencia"]["fuente_valor_minimo"] = min_smmlv["fuente"]

    # 5b. Valor mínimo de experiencia (% del presupuesto oficial)
    if not min_smmlv:
        porcentajes = _buscar_valor_porcentaje(
            texto_pliego,
            [
                "experiencia general",
                "experiencia especifica",
                "experiencia",
                "valor del contrato",
            ],
            palabras_excluir=[
                "anticipo",
                "pago anticipado",
                "capital de trabajo",
                "aiu",
                "administracion",
                "imprevisto",
                "utilidad",
            ],
        )
        # Filtrar porcentajes plausibles (10% - 100%)
        porcentajes_plausibles = [p for p in porcentajes if 10 <= p["valor"] <= 100]
        if porcentajes_plausibles:
            # Tomar el menor valor mínimo asociado a experiencia/presupuesto
            min_pct = min(porcentajes_plausibles, key=lambda x: x["valor"])
            resultado["experiencia"]["valor_minimo_po_pct"] = min_pct["valor"]
            resultado["experiencia"]["valor_minimo_cop"] = round(presupuesto * min_pct["valor"] / 100)

    # 6. Buscar Matriz 1 Experiencia y cruzar con actividad principal
    matriz1_path = _encontrar_matriz1(documentos_proceso)
    if matriz1_path and resultado.get("actividad_principal"):
        actividad_str = resultado["actividad_principal"].get("descripcion", "")
        matriz = _extraer_actividad_matriz1(matriz1_path, actividad_str)
        if matriz:
            resultado["experiencia"]["matriz1"] = matriz
        else:
            # Intentar con código + descripción corta
            codigo = resultado["actividad_principal"].get("codigo", "")
            matriz = _extraer_actividad_matriz1(matriz1_path, codigo)
            if matriz:
                resultado["experiencia"]["matriz1"] = matriz

    # 7. Capacidad financiera
    if "patrimonio liquido" in texto_norm or "patrimonio" in texto_norm:
        pct = _buscar_valor_porcentaje(texto_pliego, ["patrimonio liquido", "patrimonio"])
        pct_validos = [p for p in pct if 1 <= p["valor"] <= 50]
        if pct_validos:
            min_p = min(pct_validos, key=lambda x: x["valor"])
            resultado["capacidad_financiera"]["patrimonio_minimo_po_pct"] = min_p["valor"]
            resultado["capacidad_financiera"]["patrimonio_minimo_cop"] = round(
                presupuesto * min_p["valor"] / 100
            )

    # Indicadores financieros del texto del pliego
    indicadores = []
    for indicador, palabras in [
        ("liquidez", ["liquidez", "razon corriente", "indicador de liquidez"]),
        ("endeudamiento", ["endeudamiento", "nivel de endeudamiento"]),
        ("cobertura", ["cobertura", "intereses"]),
        ("rentabilidad", ["rentabilidad", "roe", "roi"]),
    ]:
        if any(p in texto_norm for p in palabras):
            indicadores.append(indicador)

    # Matriz 2 — Indicadores financieros y organizacionales (valores concertados)
    matriz2_path = _encontrar_matriz2(documentos_proceso)
    matriz2_data = _extraer_matriz2(matriz2_path) if matriz2_path else None
    if matriz2_data:
        resultado["capacidad_financiera"]["matriz2"] = matriz2_data
        # Complementar/actualizar indicadores requeridos con los de la Matriz 2
        for cat in matriz2_data.get("categorias", []):
            if cat not in indicadores and cat != "capital_trabajo":
                indicadores.append(cat)
        if "capital_trabajo" in matriz2_data.get("categorias", []):
            resultado["capacidad_financiera"]["capital_trabajo_requerido"] = True

    if indicadores:
        resultado["capacidad_financiera"]["indicadores_requeridos"] = indicadores

    # 8. Capacidad residual
    resultado["capacidad_residual"] = _extraer_capacidad_residual(texto_pliego)

    # 9. Documentos requeridos (habilitantes y de oferta)
    documentos_map = {
        "rup": ["registro unico de proponentes", "registro único de proponentes", "rup"],
        "autorizacion_datos_personales": [
            "autorizacion de datos personales",
            "autorización de datos personales",
            "autorizacion para el tratamiento de datos personales",
            "autorización para el tratamiento de datos personales",
            "tratamiento de datos personales",
            "datos personales",
            "ley 1581 de 2012",
            "formato 11",
            "formato11",
        ],
        "estados_financieros": [
            "estados financieros",
            "estado de situacion financiera",
            "estado de situación financiera",
            "estado de resultados",
            "balance general",
            "estados financieros auditados",
            "estados contables",
        ],
        "certificados_experiencia": [
            "certificados de experiencia",
            "certificado de experiencia",
            "experiencia del proponente",
            "soportes de experiencia",
            "actas de liquidacion",
            "actas de liquidación",
            "formato 3",
        ],
        "paz_salvo_parafiscales": ["paz y salvo", "parafiscales", "pago de aportes", "seguridad social"],
        "poliza_seriedad": [
            "poliza de seriedad",
            "póliza de seriedad",
            "seriedad de la oferta",
            "garantia de seriedad",
            "garantia de seriedad de la oferta",
            "garantia de oferta",
            "caucion de seriedad",
            "caucion de oferta",
        ],
        "propuesta_tecnica": [
            "propuesta tecnica",
            "propuesta técnica",
            "oferta tecnica",
            "oferta técnica",
            "plan de trabajo",
            "metodologia",
            "metodología",
            "programa de trabajo",
        ],
        "propuesta_economica": [
            "propuesta economica",
            "propuesta económica",
            "oferta economica",
            "oferta económica",
            "formato de precios",
            "precios unitarios",
        ],
        "carta_presentacion": ["carta de presentacion", "carta de presentación", "presentacion de la oferta"],
        # Matrices y formatos técnicos/financieros que el proponente debe diligenciar
        "matriz1_experiencia": ["matriz 1", "matriz1", "experiencia requerida"],
        "matriz2_indicadores": ["matriz 2", "matriz2", "indicadores financieros"],
        "matriz3_riesgos": ["matriz 3", "matriz3", "riesgos"],
        "capacidad_financiera": ["capacidad financiera", "formato 4", "formato4"],
        "capacidad_residual": ["capacidad residual", "formato 5", "formato5"],
        "bienes_relevantes": ["matriz 4", "matriz4", "bienes relevantes"],
    }
    for doc_id, palabras in documentos_map.items():
        if any(p in texto_norm for p in palabras):
            resultado["documentos_requeridos"].append(doc_id)

    # 10. Factores de calidad / puntaje
    factores = {}
    if "factor de calidad" in texto_norm:
        factores["factor_calidad"] = True
    if "industria nacional" in texto_norm:
        factores["industria_nacional"] = True
    if "mipyme" in texto_norm:
        factores["mipyme"] = True
    if "empresas de mujeres" in texto_norm or "emprendimiento" in texto_norm:
        factores["empresas_mujeres"] = True
    if factores:
        resultado["factores_calidad"] = factores

    # Advertencias / notas
    if not matriz1_path:
        resultado["advertencias"].append("No se encontró Matriz 1 — Experiencia entre los documentos.")
    if not matriz2_path:
        resultado["advertencias"].append("No se encontró Matriz 2 — Indicadores Financieros entre los documentos.")
    if not resultado["experiencia"].get("valor_minimo_po_pct") and not resultado["experiencia"].get("valor_minimo_smmlv"):
        resultado["advertencias"].append(
            "No se pudo determinar un valor mínimo de experiencia en % del presupuesto oficial ni en SMMLV."
        )
    if resultado["capacidad_residual"].get("requerida") and not resultado["capacidad_residual"].get("formula_crpc_corto_plazo"):
        resultado["advertencias"].append(
            "Se requiere capacidad residual pero no se pudo extraer la fórmula de CRPC del texto."
        )

    return resultado


def resumen_requisitos_para_cliente(requisitos: dict[str, Any]) -> list[dict]:
    """Convierte los requisitos estructurados en una lista legible de faltantes/verificaciones."""
    items = []
    exp = requisitos.get("experiencia", {})
    if exp.get("min_contratos"):
        items.append({
            "campo": "Experiencia mínima",
            "requerido": f"{exp['min_contratos']} contrato(s)",
            "detalle": exp.get("tipos_obra", []),
        })
    if exp.get("valor_minimo_po_pct"):
        items.append({
            "campo": "Valor mínimo de experiencia",
            "requerido": f"{exp['valor_minimo_po_pct']}% del presupuesto oficial",
            "detalle": f"${exp.get('valor_minimo_cop', 0):,} COP",
        })
    if exp.get("matriz1"):
        m = exp["matriz1"]
        items.append({
            "campo": "Matriz 1 — Experiencia",
            "requerido": m.get("actividad"),
            "detalle": {
                "general": m.get("experiencia_general", "")[:200],
                "especifica": m.get("experiencia_especifica", "")[:200],
            },
        })

    cf = requisitos.get("capacidad_financiera", {})
    if cf.get("patrimonio_minimo_po_pct"):
        items.append({
            "campo": "Patrimonio líquido mínimo",
            "requerido": f"{cf['patrimonio_minimo_po_pct']}% del presupuesto oficial",
            "detalle": f"${cf.get('patrimonio_minimo_cop', 0):,} COP",
        })
    if cf.get("indicadores_requeridos"):
        items.append({
            "campo": "Indicadores financieros",
            "requerido": cf["indicadores_requeridos"],
        })
    if cf.get("matriz2"):
        items.append({
            "campo": "Matriz 2 — Indicadores financieros",
            "requerido": cf["matriz2"].get("resumen", {}),
        })

    cr = requisitos.get("capacidad_residual", {})
    if cr.get("requerida"):
        detalle = []
        if cr.get("formula_crpc_corto_plazo"):
            detalle.append(cr["formula_crpc_corto_plazo"])
        if cr.get("formula_crpc_largo_plazo"):
            detalle.append(cr["formula_crpc_largo_plazo"])
        if cr.get("requisito_crp_crpc"):
            detalle.append(cr["requisito_crp_crpc"])
        if cr.get("min_crp_pct"):
            detalle.append(f"CRP mínimo: {cr['min_crp_pct']}%")
        items.append({
            "campo": "Capacidad residual",
            "requerido": "Requerida",
            "detalle": " | ".join(detalle) if detalle else "Métodología en pliego",
        })

    return items
